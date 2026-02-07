"""
Document loaders for various file formats.
"""
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Iterator
import io


class BaseLoader(ABC):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def load(self, content: bytes) -> str:
        """Load and extract text from document."""
        pass


class PDFLoader(BaseLoader):
    """PDF document loader using pypdf."""
    
    def load(self, content: bytes) -> str:
        """Extract text from PDF."""
        from pypdf import PdfReader
        
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)


class DocxLoader(BaseLoader):
    """Microsoft Word document loader."""
    
    def load(self, content: bytes) -> str:
        """Extract text from DOCX."""
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)


class TextLoader(BaseLoader):
    """Plain text file loader."""
    
    def load(self, content: bytes) -> str:
        """Decode text content."""
        # Try UTF-8 first, then fall back to other encodings
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort: decode with errors ignored
        return content.decode("utf-8", errors="ignore")


class MarkdownLoader(TextLoader):
    """Markdown file loader with basic cleanup."""
    
    def load(self, content: bytes) -> str:
        """Load markdown and optionally convert to plain text."""
        text = super().load(content)
        
        # Basic markdown cleanup (keep structure but clean up)
        # In production, use a proper markdown parser
        lines = []
        for line in text.split("\n"):
            # Remove HTML comments
            if "<!--" in line and "-->" in line:
                continue
            lines.append(line)
        
        return "\n".join(lines)


class HTMLLoader(BaseLoader):
    """HTML document loader using BeautifulSoup."""
    
    def load(self, content: bytes) -> str:
        """Extract text from HTML."""
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(content, "html.parser")
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "header", "footer"]):
            script.decompose()
        
        # Get text
        text = soup.get_text(separator="\n")
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        
        return text


def get_loader(file_type: str) -> BaseLoader:
    """
    Get the appropriate loader for a file type.
    
    Args:
        file_type: File extension (e.g., ".pdf", ".docx").
        
    Returns:
        A loader instance.
    """
    loaders = {
        ".pdf": PDFLoader,
        ".docx": DocxLoader,
        ".txt": TextLoader,
        ".md": MarkdownLoader,
        ".html": HTMLLoader,
        ".htm": HTMLLoader,
    }
    
    if file_type not in loaders:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    return loaders[file_type]()
